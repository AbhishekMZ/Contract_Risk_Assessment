# src/contract_analysis/imp_pipeline.py - Improved contract analysis pipeline
import os
import json
import re
import pandas as pd
from typing import Dict, List, Any, Union, Tuple
from pathlib import Path

# NLP libraries
import nltk
import spacy
from transformers import BertTokenizer, BertForSequenceClassification, pipeline

# File processing libraries
import pdfplumber  # Enhanced PDF extraction
import docx  # Enhanced DOCX extraction

# Smart contract analysis
from web3 import Web3
from solidity_parser import parser


class ContractParser:
    """Base class for contract parsing with common functionality"""
    
    def __init__(self):
        self.contract_text = None
        self.contract_metadata = {}
        self.parsed_sections = {}
        
    def load_contract(self, file_path: str) -> bool:
        """Load contract from file"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if not os.path.exists(file_path):
                print(f"File not found: {file_path}")
                return False
                
            self.contract_metadata['file_path'] = file_path
            self.contract_metadata['file_type'] = file_ext
            self.contract_metadata['file_name'] = os.path.basename(file_path)
            self.contract_metadata['file_size'] = os.path.getsize(file_path)
            
            return True
        except Exception as e:
            print(f"Error loading contract: {str(e)}")
            return False
    
    def extract_metadata(self) -> Dict:
        """Extract metadata from the contract"""
        return self.contract_metadata
        
    def get_parsed_sections(self) -> Dict:
        """Return parsed contract sections"""
        return self.parsed_sections


class LegalContractParser(ContractParser):
    """Parser for traditional legal contracts (PDF, DOC, etc.)"""
    
    def __init__(self):
        super().__init__()
        # Initialize NLP components
        nltk.download('punkt', quiet=True)
        nltk.download('stopwords', quiet=True)
        try:
            self.nlp = spacy.load('en_core_web_sm')
        except OSError:
            print("Downloading spacy model...")
            spacy.cli.download('en_core_web_sm')
            self.nlp = spacy.load('en_core_web_sm')
        
    def load_contract(self, file_path: str) -> bool:
        """Load and extract text from legal document"""
        if not super().load_contract(file_path):
            return False
            
        file_ext = self.contract_metadata['file_type']
        
        try:
            if file_ext == '.pdf':
                self._extract_from_pdf(file_path)
            elif file_ext in ['.doc', '.docx']:
                self._extract_from_word(file_path)
            elif file_ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    self.contract_text = f.read()
            else:
                print(f"Unsupported file type: {file_ext}")
                return False
            
            # Add text length to metadata
            if self.contract_text:
                self.contract_metadata['text_length'] = len(self.contract_text)
                return True
            else:
                print(f"Warning: No text extracted from {file_path}")
                return False
                
        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
            return False
    
    def _extract_from_pdf(self, file_path: str) -> None:
        """Extract text from PDF file using pdfplumber with enhanced error handling."""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                # Add page count to metadata
                self.contract_metadata['page_count'] = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            self.contract_text = text
            if not self.contract_text:
                print(f"Warning: No text extracted from PDF: {file_path}")
        except Exception as e:
            print(f"Error extracting text from PDF {file_path}: {str(e)}")
            self.contract_text = ""  # Ensure contract_text is set even on error
    
    def _extract_from_word(self, file_path: str) -> None:
        """Extract text from Word document using python-docx with enhanced extraction."""
        try:
            doc = docx.Document(file_path)
            
            # Add metadata
            self.contract_metadata['paragraph_count'] = len(doc.paragraphs)
            
            # Extract text from paragraphs
            full_text = [para.text for para in doc.paragraphs]
            
            # Extract tables if present
            for table in doc.tables:
                for row in table.rows:
                    table_row = []
                    for cell in row.cells:
                        table_row.append(cell.text)
                    full_text.append(" | ".join(table_row))
            
            self.contract_text = "\n".join(full_text)
            
            if not self.contract_text:
                print(f"Warning: No text extracted from DOCX: {file_path}")
                
        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {str(e)}")
            self.contract_text = ""  # Ensure contract_text is set even on error
    
    def parse_sections(self) -> Dict:
        """Parse contract into logical sections using improved heuristics."""
        if not self.contract_text:
            return {}
        
        # Use spaCy to process the document
        doc = self.nlp(self.contract_text[:1000000])  # Limit to prevent memory issues with large docs
        
        sections = {}
        current_section_title = "Introduction / Preamble"
        current_section_content = []

        # More robust section detection using regex patterns
        for para_text in self.contract_text.split('\n\n'):  # Double newline separates paragraphs
            para_text = para_text.strip()
            if not para_text:
                continue

            # Enhanced heuristics for section headers
            is_header = False
            
            # Check for common header patterns:
            # 1. All uppercase text with less than 10 words
            # 2. Text starting with "Article", "Section", etc.
            # 3. Numbered headers like "1.", "1.1", "I.", etc.
            if (
                (len(para_text.split()) < 10 and para_text.isupper()) or
                re.match(r"^(ARTICLE|SECTION|CLAUSE)(\s+[IVX\d]+\.?)?", para_text, re.IGNORECASE) or
                re.match(r"^(\d+\.)+\d*\s+\w+", para_text) or
                re.match(r"^[IVX]+\.\s+\w+", para_text)
            ):
                is_header = True

            if is_header:
                if current_section_content:  # Save previous section
                    sections[current_section_title] = "\n".join(current_section_content)
                current_section_title = para_text
                current_section_content = []
            else:
                current_section_content.append(para_text)

        if current_section_content:  # Save the last section
            sections[current_section_title] = "\n".join(current_section_content)

        self.parsed_sections = sections
        
        # Add section count to metadata
        self.contract_metadata['section_count'] = len(sections)
        
        return sections
    
    def extract_entities(self) -> Dict:
        """Extract named entities from the contract with improved categorization."""
        entities = {
            'organizations': [],
            'people': [],
            'dates': [],
            'money_amounts': [],
            'locations': [],
            'percentages': [],
            'products': [],
            'laws': []
        }
        
        if not self.contract_text:
            return entities
        
        # Process the document in chunks to avoid memory issues
        chunk_size = 100000  # Process in ~100k character chunks
        chunks = [self.contract_text[i:i + chunk_size] 
                 for i in range(0, len(self.contract_text), chunk_size)]
        
        for chunk in chunks:
            doc = self.nlp(chunk)
            
            for ent in doc.ents:
                if ent.label_ == 'ORG':
                    entities['organizations'].append(ent.text)
                elif ent.label_ == 'PERSON':
                    entities['people'].append(ent.text)
                elif ent.label_ == 'DATE':
                    entities['dates'].append(ent.text)
                elif ent.label_ == 'MONEY':
                    entities['money_amounts'].append(ent.text)
                elif ent.label_ == 'GPE' or ent.label_ == 'LOC':
                    entities['locations'].append(ent.text)
                elif ent.label_ == 'PERCENT':
                    entities['percentages'].append(ent.text)
                elif ent.label_ == 'PRODUCT':
                    entities['products'].append(ent.text)
                elif ent.label_ == 'LAW':
                    entities['laws'].append(ent.text)
        
        # Remove duplicates and sort alphabetically
        for key in entities:
            entities[key] = sorted(list(set(entities[key])))
            
        return entities


class SmartContractParser(ContractParser):
    """Parser for blockchain smart contracts (Solidity)"""
    
    def __init__(self):
        super().__init__()
        
    def load_contract(self, file_path: str) -> bool:
        """Load smart contract from file"""
        if not super().load_contract(file_path):
            return False
            
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                self.contract_text = f.read()
            
            self.contract_metadata['file_size'] = os.path.getsize(file_path)
            return True
        except Exception as e:
            print(f"Error loading smart contract: {str(e)}")
            return False
    
    def parse_contract(self) -> Dict:
        """Parse Solidity contract structure"""
        if not self.contract_text:
            return {}
            
        try:
            # Parse Solidity code using solidity-parser
            parsed_data = parser.parse(self.contract_text)
            
            # Extract key elements
            contracts = []
            functions = []
            events = []
            modifiers = []
            state_variables = []
            
            # Process AST to extract contract elements
            for node in parsed_data['children']:
                if node['type'] == 'ContractDefinition':
                    contract_name = node['name']
                    contract_kind = node.get('kind', 'contract')
                    contracts.append({
                        'name': contract_name,
                        'type': contract_kind
                    })
                    
                    # Process contract body
                    for body_node in node.get('subNodes', []):
                        if body_node['type'] == 'FunctionDefinition':
                            function_name = body_node.get('name', '')
                            if function_name == '':
                                if body_node.get('isConstructor', False):
                                    function_name = 'constructor'
                                else:
                                    function_name = 'fallback'
                                    
                            functions.append({
                                'contract': contract_name,
                                'name': function_name,
                                'visibility': body_node.get('visibility', 'internal'),
                                'is_payable': body_node.get('stateMutability', '') == 'payable',
                                'is_view': body_node.get('stateMutability', '') in ['view', 'pure']
                            })
                        
                        elif body_node['type'] == 'EventDefinition':
                            events.append({
                                'contract': contract_name,
                                'name': body_node['name']
                            })
                            
                        elif body_node['type'] == 'ModifierDefinition':
                            modifiers.append({
                                'contract': contract_name,
                                'name': body_node['name']
                            })
                            
                        elif body_node['type'] == 'StateVariableDeclaration':
                            for variable in body_node.get('variables', []):
                                state_variables.append({
                                    'contract': contract_name,
                                    'name': variable['name'],
                                    'type': variable.get('typeName', {}).get('name', 'unknown')
                                })
            
            # Store the parsed data in sections
            self.parsed_sections = {
                'contracts': contracts,
                'functions': functions,
                'events': events,
                'modifiers': modifiers,
                'state_variables': state_variables
            }
            
            return self.parsed_sections
            
        except Exception as e:
            print(f"Error parsing Solidity contract: {str(e)}")
            return {}
    
    def identify_security_patterns(self) -> Dict:
        """Identify common security patterns and potential issues"""
        issues = []
        
        if not self.contract_text:
            return {'issues': issues}
            
        # Check for reentrancy vulnerabilities
        if 'call.value' in self.contract_text:
            issues.append({
                'type': 'reentrancy',
                'severity': 'high',
                'description': 'Potential reentrancy vulnerability detected (call.value)'
            })
            
        # Check for unchecked external calls
        if '.call(' in self.contract_text and 'require(' not in self.contract_text:
            issues.append({
                'type': 'unchecked_call',
                'severity': 'medium',
                'description': 'External call without checking return value'
            })
            
        # Check for tx.origin usage
        if 'tx.origin' in self.contract_text:
            issues.append({
                'type': 'tx_origin',
                'severity': 'medium',
                'description': 'Use of tx.origin for authentication'
            })
            
        # Check for block timestamp dependence
        if 'block.timestamp' in self.contract_text or 'now' in self.contract_text:
            issues.append({
                'type': 'timestamp_dependence',
                'severity': 'low',
                'description': 'Dependence on block timestamp'
            })
            
        return {'issues': issues}


class ContractAnalysisPipeline:
    """Pipeline to analyze contracts and generate reports"""
    
    def __init__(self):
        self.parsers = {
            'legal': LegalContractParser(),
            'smart': SmartContractParser()
        }
        
        # Initialize BERT model for contract classification
        self.bert_tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
        self.bert_model = BertForSequenceClassification.from_pretrained('bert-base-uncased')
        self.text_classifier = pipeline('text-classification', model=self.bert_model, tokenizer=self.bert_tokenizer)
        
    def analyze_contract(self, file_path: str) -> Dict:
        """Analyze a contract file and return structured results"""
        # Determine contract type based on extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext in ['.pdf', '.doc', '.docx', '.txt', '.md']:
            parser = self.parsers['legal']
            contract_type = 'legal'
        elif file_ext in ['.sol']:
            parser = self.parsers['smart']
            contract_type = 'smart'
        else:
            return {'error': f'Unsupported file type: {file_ext}'}
        
        # Load and parse the contract
        if not parser.load_contract(file_path):
            return {'error': 'Failed to load contract'}
            
        # Extract contract metadata
        metadata = parser.extract_metadata()
        
        # Perform type-specific analysis
        analysis_results = {}
        
        if contract_type == 'legal':
            # Parse sections
            sections = parser.parse_sections()
            
            # Extract entities
            entities = parser.extract_entities()
            
            # Classify contract type if text is available
            contract_category = 'unknown'
            if parser.contract_text and len(parser.contract_text) > 20:
                sample = parser.contract_text[:512]  # Use beginning for classification
                try:
                    classification = self.text_classifier(sample)
                    contract_category = classification[0]['label']
                except Exception as e:
                    print(f"Error during contract classification: {str(e)}")
            
            analysis_results = {
                'sections': sections,
                'entities': entities,
                'category': contract_category
            }
            
        elif contract_type == 'smart':
            # Parse contract structure
            structure = parser.parse_contract()
            
            # Identify security patterns
            security = parser.identify_security_patterns()
            
            analysis_results = {
                'structure': structure,
                'security': security
            }
        
        # Return combined results
        return {
            'metadata': metadata,
            'type': contract_type,
            'analysis': analysis_results
        }
